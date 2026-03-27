import os
import secrets
import sys
from pathlib import Path
from datetime import datetime, timedelta
from typing import List
from io import BytesIO
import httpx
from db import SessionLocal

from fastapi import UploadFile, File
from fastapi import FastAPI, Depends, HTTPException, Request
from fastapi.responses import HTMLResponse, RedirectResponse, Response, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session, joinedload
from passlib.context import CryptContext
# from jose import jwt, JWTError

from db import Base, engine, get_db
from models import User, Calculation, CalculationYear, RiskLevel, UserSession
import schemas

# ---------- DOCX / Charts ----------
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Cm

# ВАЖНО: без GUI-бэкенд для matplotlib
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# ---------- Excel ----------
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage
from openpyxl import load_workbook
from PIL import Image as PILImage
import re

# ---------- PDF ----------
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
)
# ---------- ПУТИ для PyInstaller ----------
if getattr(sys, 'frozen', False):
    BASE_DIR = Path(sys._MEIPASS)
else:
    BASE_DIR = Path(__file__).parent

# # --- Пути к шаблонам/статикам (АБСОЛЮТНЫЕ!) ---
# BASE_DIR = Path(__file__).parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

app = FastAPI(title="Investment Calculator")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"], allow_credentials=True)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

Base.metadata.create_all(bind=engine)

def seed_risk_levels():
    db = SessionLocal()
    try:
        if db.query(RiskLevel).count() == 0:
            db.add_all([
                RiskLevel(name="Низкий риск", expected_return=5.0, description="Минимальный риск, 3–7% годовых"),
                RiskLevel(name="Средний риск", expected_return=9.5, description="Умеренный риск, 7–12% годовых"),
                RiskLevel(name="Высокий риск", expected_return=15.0, description="Высокий риск, от 12% годовых"),
            ])
            db.commit()
    finally:
        db.close()

seed_risk_levels()
# SECRET_KEY = os.getenv("SECRET_KEY", "dev_secret_change_me")
# ALGO = "HS256"
pwd = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")


# =================== helpers ===================

# def token_for(user_id: int) -> str:
#     return jwt.encode({"sub": str(user_id), "exp": datetime.utcnow() + timedelta(days=7)}, SECRET_KEY, algorithm=ALGO)

def get_current_user(request: Request, db: Session = Depends(get_db)):
    token = request.cookies.get("session_token")

    if not token:
        return None

    session = db.query(UserSession).filter(
        UserSession.token == token
    ).first()

    if not session:
        return None

    if session.expires_at < datetime.utcnow():
        db.delete(session)
        db.commit()
        return None

    return session.user

def require_user(user: User = Depends(get_current_user)):
    if not user:
        raise HTTPException(401, "Не авторизован")
    return user
# def current_user(request: Request, db: Session) -> User | None:
#     tok = request.cookies.get("access_token")
#     if not tok:
#         return None
#     try:
#         uid = int(jwt.decode(tok, SECRET_KEY, algorithms=[ALGO]).get("sub"))
#     except JWTError:
#         return None
#     return db.query(User).filter_by(id=uid).first()

def create_session(db: Session, user_id: int):
    token = secrets.token_hex(32)

    session = UserSession(
        user_id=user_id,
        token=token,
        expires_at=datetime.utcnow() + timedelta(days=7)
    )

    db.add(session)
    db.commit()

    return token

def require_user(request: Request, db: Session = Depends(get_db)) -> User:
    u = get_current_user(request, db)
    if not u:
        raise HTTPException(401, "Not authenticated")
    return u

def rub(x: float) -> str:
    return f"{x:,.0f} ₽".replace(",", " ")

def _chart_png_bytes(years: list[int], totals: list[float]) -> BytesIO:
    """Общий генератор PNG-графика (без GUI)."""
    fig, ax = plt.subplots(figsize=(6.5, 3))
    ax.plot(years, totals, marker="o")
    ax.set_xlabel("Год")
    ax.set_ylabel("Сумма, ₽")
    ax.grid(True, alpha=0.3)
    buf = BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=200)
    plt.close(fig)
    buf.seek(0)
    return buf


# =================== pages ===================

@app.get("/", response_class=HTMLResponse)
def home(request: Request, db: Session = Depends(get_db)):
    return templates.TemplateResponse("home.html", {"request": request, "user": get_current_user(request, db)})

@app.get("/calc", response_class=HTMLResponse)
def calc_page(request: Request, db: Session = Depends(get_db)):
    return templates.TemplateResponse("calc.html", {"request": request, "user": get_current_user(request, db)})

@app.get("/profile", response_class=HTMLResponse)
def profile_page(request: Request, db: Session = Depends(get_db)):
    return templates.TemplateResponse("profile.html", {"request": request, "user": get_current_user(request, db)})


# =================== auth api ===================

@app.post("/auth/register")
async def register(request: Request, db: Session = Depends(get_db)):
    ct = request.headers.get("content-type", "")
    try:
        data = await (request.json() if "application/json" in ct else (await request.form()))
        data = dict(data)
    except Exception:
        data = {}

    name = (data.get("name") or "").strip()
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""
    if not (name and email and password):
        raise HTTPException(422, "name, email, password обязательны")
    if db.query(User).filter_by(email=email).first():
        raise HTTPException(400, "Email уже зарегистрирован")

    u = User(name=name, email=email, password_hash=pwd.hash(password))
    db.add(u); db.commit(); db.refresh(u)
    return {"id": u.id, "name": u.name, "email": u.email}

@app.post("/auth/login")
async def login(request: Request, response: Response, db: Session = Depends(get_db)):
    data = await request.json()

    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    if not (email and password):
        raise HTTPException(422, "email и password обязательны")

    user = db.query(User).filter_by(email=email).first()

    if not user or not pwd.verify(password, user.password_hash):
        raise HTTPException(401, "Неверный email или пароль")

    token = create_session(db, user.id)

    response = JSONResponse({
        "ok": True,
        "user": {"id": user.id, "name": user.name, "email": user.email}
    })

    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        samesite="lax",
        max_age=7 * 24 * 3600
    )

    return response

    token = token_for(u.id)
    if "application/json" in ct:
        resp = JSONResponse({"ok": True, "user": {"id": u.id, "name": u.name, "email": u.email}})
    else:
        resp = RedirectResponse(url="/", status_code=302)
    resp.set_cookie("access_token", token, httponly=True, samesite="lax", max_age=7*24*3600)
    return resp

@app.post("/auth/logout")
def logout(request: Request, response: Response, db: Session = Depends(get_db)):
    token = request.cookies.get("session_token")

    if token:
        db.query(UserSession).filter(
            UserSession.token == token
        ).delete()
        db.commit()

    resp = RedirectResponse(url="/", status_code=302)
    resp.delete_cookie("session_token")

    return resp


# =================== calc core ===================

def run_calc(inp: schemas.CalcInput) -> schemas.CalcResult:
    P = float(inp.initial_amount)
    m = float(inp.monthly_contribution)
    r = float(inp.annual_rate) / 100.0
    years = int(inp.years)

    bal = P
    contrib = P
    schedule = []
    for y in range(1, years + 1):
        for _ in range(12):
            bal *= (1 + r / 12.0)
            bal += m
            contrib += m
        schedule.append(
            schemas.YearRow(
                year=y,
                total=round(bal, 2),
                contributions=round(contrib, 2),
                profit=round(bal - contrib, 2),
            )
        )
    return schemas.CalcResult(
        final_amount=round(bal, 2),
        total_contributions=round(contrib, 2),
        profit=round(bal - contrib, 2),
        schedule=schedule,
    )

@app.post("/api/calc/run", response_model=schemas.CalcResult)
def api_run(inp: schemas.CalcInput):
    return run_calc(inp)

@app.post("/api/calc/save", response_model=schemas.CalcOut)
def api_save(p: schemas.CalcCreate, u: User = Depends(require_user), db: Session = Depends(get_db)):
    res = run_calc(p)

    calc = Calculation(
        user_id=u.id,
        title=p.title.strip(),
        initial_amount=p.initial_amount,
        monthly_contribution=p.monthly_contribution,
        annual_rate=p.annual_rate,
        years=p.years,
        final_amount=res.final_amount,
        total_contributions=res.total_contributions,
        profit=res.profit,
    )
    db.add(calc)
    db.flush()  # получим calc.id

    rows = [
        CalculationYear(
            calculation_id=calc.id,
            year=r.year,
            total=r.total,
            contributions=r.contributions,
            profit=r.profit,
        )
        for r in res.schedule
    ]
    db.add_all(rows)
    db.commit()
    db.refresh(calc)

    return schemas.CalcOut(
        id=calc.id,
        title=calc.title,
        initial_amount=calc.initial_amount,
        monthly_contribution=calc.monthly_contribution,
        annual_rate=calc.annual_rate,
        years=calc.years,
        final_amount=calc.final_amount,
        total_contributions=calc.total_contributions,
        profit=calc.profit,
        schedule=[
            schemas.YearRow(year=y.year, total=y.total, contributions=y.contributions, profit=y.profit)
            for y in calc.years_rows
        ],
    )

@app.get("/api/calc/list", response_model=List[schemas.CalcOut])
def api_list(u: User = Depends(require_user), db: Session = Depends(get_db)):
    items = (
        db.query(Calculation)
        .options(joinedload(Calculation.years_rows))
        .filter(Calculation.user_id == u.id)
        .order_by(Calculation.created_at.desc())
        .all()
    )
    out: List[schemas.CalcOut] = []
    for i in items:
        out.append(
            schemas.CalcOut(
                id=i.id,
                title=i.title,
                initial_amount=i.initial_amount,
                monthly_contribution=i.monthly_contribution,
                annual_rate=i.annual_rate,
                years=i.years,
                final_amount=i.final_amount,
                total_contributions=i.total_contributions,
                profit=i.profit,
                schedule=[
                    schemas.YearRow(year=y.year, total=y.total, contributions=y.contributions, profit=y.profit)
                    for y in i.years_rows
                ],
            )
        )
    return out


# =================== reports ===================

# ---- WORD (заменяет старый CSV) ----
@app.get("/api/report/{calc_id}")
def api_report_word(calc_id: int, u: User = Depends(require_user), db: Session = Depends(get_db)):
    calc = (
        db.query(Calculation)
        .options(joinedload(Calculation.years_rows))
        .filter(Calculation.id == calc_id, Calculation.user_id == u.id)
        .first()
    )
    if not calc:
        raise HTTPException(404, "Расчёт не найден")

    years = [r.year for r in calc.years_rows]
    totals = [r.total for r in calc.years_rows]
    png_buf = _chart_png_bytes(years, totals)

    tpl_path = BASE_DIR / "templates" / "docx" / "calc_report.docx"
    if not tpl_path.exists():
        raise HTTPException(500, f"Шаблон не найден: {tpl_path}")

    doc = DocxTemplate(str(tpl_path))
    context = {
        "title": calc.title,
        "user_name": u.name,
        "user_email": u.email,
        "created_at": calc.created_at.strftime("%d.%m.%Y %H:%M"),
        "final_amount": rub(calc.final_amount),
        "total_contributions": rub(calc.total_contributions),
        "profit": rub(calc.profit),
        "chart_image": InlineImage(doc, png_buf, width=Cm(16)),
    }
    doc.render(context)

    # первая таблица — для данных по годам
    table = doc.tables[0]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)
    for r in calc.years_rows:
        row = table.add_row()
        row.cells[0].text = str(r.year)
        row.cells[1].text = rub(r.total)
        row.cells[2].text = rub(r.contributions)
        row.cells[3].text = rub(r.profit)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename=calc_report_{calc_id}.docx'}
    )


# ---- EXCEL (.xlsx) ----
@app.get("/api/report/excel/{calc_id}")
def api_report_excel(calc_id: int, u: User = Depends(require_user), db: Session = Depends(get_db)):
    calc = (
        db.query(Calculation)
        .options(joinedload(Calculation.years_rows))
        .filter(Calculation.id == calc_id, Calculation.user_id == u.id)
        .first()
    )
    if not calc:
        raise HTTPException(404, "Расчёт не найден")

    wb = Workbook()
    ws = wb.active
    ws.title = "Расчёт"

    ws["A1"] = f"Отчёт по расчёту: {calc.title}"
    ws["A2"] = f"Пользователь: {u.name} ({u.email})"
    ws["A3"] = f"Итоговая сумма: {rub(calc.final_amount)}"
    ws["A4"] = f"Общие взносы: {rub(calc.total_contributions)}"
    ws["A5"] = f"Доход: {rub(calc.profit)}"

    start_row = 7
    ws.cell(row=start_row, column=1, value="Год")
    ws.cell(row=start_row, column=2, value="Сумма")
    ws.cell(row=start_row, column=3, value="Взносы")
    ws.cell(row=start_row, column=4, value="Доходность")

    r = start_row + 1
    for y in calc.years_rows:
        ws.cell(row=r, column=1, value=y.year)
        ws.cell(row=r, column=2, value=y.total)
        ws.cell(row=r, column=3, value=y.contributions)
        ws.cell(row=r, column=4, value=y.profit)
        r += 1

    # автоширина
    for col in range(1, 5):
        letter = get_column_letter(col)
        maxlen = 10
        for cell in ws[letter]:
            maxlen = max(maxlen, len(str(cell.value)) if cell.value is not None else 0)
        ws.column_dimensions[letter].width = maxlen + 2

    # график как PNG
    years = [row.year for row in calc.years_rows]
    totals = [row.total for row in calc.years_rows]
    png_buf = _chart_png_bytes(years, totals)
    pil_img = PILImage.open(png_buf)
    tmp_buf = BytesIO()
    pil_img.save(tmp_buf, format="PNG")
    tmp_buf.seek(0)
    img = XLImage(tmp_buf)
    img.anchor = "E1"
    ws.add_image(img)

    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename=calc_report_{calc_id}.xlsx'}
    )


# ---- PDF ----
@app.get("/api/report/pdf/{calc_id}")
def api_report_pdf(calc_id: int, u: User = Depends(require_user), db: Session = Depends(get_db)):
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    import traceback

    # 1) данные
    calc = (
        db.query(Calculation)
        .options(joinedload(Calculation.years_rows))
        .filter(Calculation.id == calc_id, Calculation.user_id == u.id)
        .first()
    )
    if not calc:
        raise HTTPException(404, "Расчёт не найден")

    years = [r.year for r in calc.years_rows]
    totals = [r.total for r in calc.years_rows]
    png_buf = _chart_png_bytes(years, totals)

    # 2) пытаемся найти шрифты: сначала в static/fonts, иначе — берем из Matplotlib
    font_dir = BASE_DIR / "static" / "fonts"
    local_reg = (font_dir / "DejaVuSans.ttf", font_dir / "DejaVuSans-Bold.ttf")

    if all(p.exists() for p in local_reg):
        regular_path, bold_path = map(str, local_reg)
    else:
        # fallback на встроенные шрифты matplotlib
        import matplotlib
        from pathlib import Path
        md = Path(matplotlib.get_data_path()) / "fonts" / "ttf"
        regular_path = str(md / "DejaVuSans.ttf")
        bold_path    = str(md / "DejaVuSans-Bold.ttf")

    try:
        pdfmetrics.registerFont(TTFont("DejaVu", regular_path))
        pdfmetrics.registerFont(TTFont("DejaVu-Bold", bold_path))
    except Exception:
        # если даже тут что-то пойдет не так — вернем понятную ошибку
        traceback.print_exc()
        raise HTTPException(500, "Не удалось зарегистрировать шрифты DejaVu для PDF")

    # 3) PDF-документ
    pdf_buf = BytesIO()
    doc = SimpleDocTemplate(
        pdf_buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36
    )

    base = getSampleStyleSheet()
    styleTitle = ParagraphStyle(name="TitleRu", parent=base["Title"], fontName="DejaVu-Bold")
    styleH3    = ParagraphStyle(name="H3Ru",    parent=base["Heading3"], fontName="DejaVu-Bold")
    styleText  = ParagraphStyle(name="BodyRu",  parent=base["Normal"], fontName="DejaVu")

    story = []
    story.append(Paragraph(f"Отчёт по расчёту: {calc.title}", styleTitle))
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"Пользователь: {u.name} ({u.email})", styleText))
    story.append(Paragraph(f"Дата: {calc.created_at.strftime('%d.%m.%Y %H:%M')}", styleText))
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"Итоговая сумма: <b>{rub(calc.final_amount)}</b>", styleText))
    story.append(Paragraph(f"Общие взносы: <b>{rub(calc.total_contributions)}</b>", styleText))
    story.append(Paragraph(f"Доход: <b>{rub(calc.profit)}</b>", styleText))
    story.append(Spacer(1, 12))

    story.append(Paragraph("График динамики", styleH3))
    story.append(Spacer(1, 6))
    story.append(RLImage(png_buf, width=520, height=240))
    story.append(Spacer(1, 12))

    data = [["Год", "Сумма", "Взносы", "Доходность"]]
    for r in calc.years_rows:
        data.append([r.year, rub(r.total), rub(r.contributions), rub(r.profit)])

    tbl = Table(data, colWidths=[60, 140, 140, 140])
    tbl.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "DejaVu"),
        ("FONTNAME", (0, 0), (-1, 0), "DejaVu-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(tbl)

    doc.build(story)
    pdf_buf.seek(0)
    return StreamingResponse(
        pdf_buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename=calc_report_{calc_id}.pdf'}
    )

    # --------------------- < БЛОК ИМПОРТА >---------------------------

def _parse_money(x) -> float:
    """Принимает int/float/str/None. Возвращает float."""
    if x is None:
        return 0.0
    if isinstance(x, (int, float)):
        return float(x)
    # строка
    s = str(x).strip()
    # убираем валюту и неразрывные пробелы, нормализуем
    s = s.replace("₽", "").replace("\u00A0", " ").replace(" ", "")
    # если европейский формат с запятой — заменим на точку
    s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        return 0.0

def _make_calc_from_rows(u: User, db: Session, title: str, rows: list[dict]) -> Calculation:
    """rows: [{'year':int,'total':float,'contributions':float,'profit':float}, ...]"""
    if not rows:
        raise HTTPException(400, "В файле не обнаружены строки таблицы")
    years = len(rows)
    final_amount = rows[-1]['total']
    total_contributions = rows[-1]['contributions']
    profit = rows[-1]['profit']

    calc = Calculation(
        user_id=u.id,
        title=title or "Импортированный расчёт",
        initial_amount=0.0,                # исходные параметры восстановить нельзя — ставим 0
        monthly_contribution=0.0,
        annual_rate=0.0,
        years=years,
        final_amount=final_amount,
        total_contributions=total_contributions,
        profit=profit,
    )
    db.add(calc); db.flush()

    for r in rows:
        db.add(CalculationYear(
            calculation_id=calc.id,
            year=int(r['year']),
            total=float(r['total']),
            contributions=float(r['contributions']),
            profit=float(r['profit']),
        ))
    db.commit(); db.refresh(calc)
    return calc

# ---------- Импорт из Word ----------
@app.post("/api/import/word", response_model=schemas.CalcOut)
async def import_word(file: UploadFile = File(...), u: User = Depends(require_user), db: Session = Depends(get_db)):
    content = await file.read()
    try:
        doc = Document(BytesIO(content))
    except Exception:
        raise HTTPException(400, "Не удалось прочитать .docx (повреждённый файл?)")

    # 1) заголовок (ищем строку "Отчёт по расчёту: <title>")
    title = file.filename.rsplit(".", 1)[0]
    for p in doc.paragraphs:
        m = re.search(r"Отч[ёе]т по расч[её]ту:\s*(.+)", p.text.strip(), re.IGNORECASE)
        if m:
            title = m.group(1).strip()
            break

    # 2) таблица: берём первую таблицу как в экспорте
    if not doc.tables:
        raise HTTPException(400, "В документе не найдена таблица с данными")
    t = doc.tables[0]
    if len(t.rows) < 2:
        raise HTTPException(400, "Таблица должна содержать заголовок и хотя бы одну строку данных")

    rows = []
    # пропускаем заголовок (row 0)
    for i in range(1, len(t.rows)):
        cells = t.rows[i].cells
        if len(cells) < 4:
            continue
        y = cells[0].text.strip()
        if not y:
            continue
        rows.append({
            "year": int(re.sub(r"\D", "", y) or "0"),
            "total": _parse_money(cells[1].text),
            "contributions": _parse_money(cells[2].text),
            "profit": _parse_money(cells[3].text),
        })

    calc = _make_calc_from_rows(u, db, title, rows)
    return schemas.CalcOut(
        id=calc.id, title=calc.title,
        initial_amount=calc.initial_amount, monthly_contribution=calc.monthly_contribution,
        annual_rate=calc.annual_rate, years=calc.years,
        final_amount=calc.final_amount, total_contributions=calc.total_contributions, profit=calc.profit,
        schedule=[schemas.YearRow(year=r.year, total=r.total, contributions=r.contributions, profit=r.profit) for r in calc.years_rows],
    )

# ---------- Импорт из Excel ----------
@app.post("/api/import/excel", response_model=schemas.CalcOut)
async def import_excel(file: UploadFile = File(...), u: User = Depends(require_user), db: Session = Depends(get_db)):
    content = await file.read()
    try:
        wb = load_workbook(BytesIO(content), data_only=True)
        ws = wb.active
    except Exception:
        raise HTTPException(400, "Не удалось прочитать .xlsx (повреждённый файл?)")

    # 1) заголовок: пробуем вытащить из A1 формата "Отчёт по расчёту: ..."
    title = file.filename.rsplit(".", 1)[0]
    a1 = (ws["A1"].value or "").strip() if ws["A1"].value else ""
    m = re.search(r"Отч[ёе]т по расч[её]ту:\s*(.+)", a1, re.IGNORECASE)
    if m:
        title = m.group(1).strip()

    # 2) найдём строку заголовка таблицы ("Год", "Сумма", "Взносы", "Доходность")
    header_row_idx = None
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True):
        if not row:
            continue
        vals = [str(v).strip() if v is not None else "" for v in row[:6]]
        if {"Год", "Сумма", "Взносы", "Доходность"}.issubset(set(vals)):
            header_row_idx = row[0].row if hasattr(row[0], "row") else None
            break

    # если не нашли через values_only (теряем адреса), сделаем повторно без values_only
    if header_row_idx is None:
        for r in ws.iter_rows(min_row=1, max_row=ws.max_row):
            vals = [ (c.value if c.value is not None else "") for c in r ]
            texts = [str(v).strip() for v in vals[:6]]
            if {"Год", "Сумма", "Взносы", "Доходность"}.issubset(set(texts)):
                header_row_idx = r[0].row
                break

    if header_row_idx is None:
        # по умолчанию предполагаем экспортный шаблон: заголовок на 7-й строке
        header_row_idx = 7

    # 3) читаем данные до первой пустой строки (числа могут быть как числа, так и отформатированные строки)
    rows = []
    r = header_row_idx + 1
    while True:
        y_cell = ws.cell(row=r, column=1).value
        t_cell = ws.cell(row=r, column=2).value
        c_cell = ws.cell(row=r, column=3).value
        p_cell = ws.cell(row=r, column=4).value

        # пустая строка — выходим
        if y_cell is None or str(y_cell).strip() == "":
            break

        # год допустим только как целое число
        try:
            year_val = int(str(y_cell).strip())
        except Exception:
            # если в A ячейке мусор — прекращаем чтение
            break

        rows.append({
            "year": year_val,
            "total": _parse_money(t_cell),
            "contributions": _parse_money(c_cell),
            "profit": _parse_money(p_cell),
        })
        r += 1

    if not rows:
        raise HTTPException(400, "Не найдено ни одной строки данных под заголовком таблицы")

    calc = _make_calc_from_rows(u, db, title, rows)
    return schemas.CalcOut(
        id=calc.id, title=calc.title,
        initial_amount=calc.initial_amount, monthly_contribution=calc.monthly_contribution,
        annual_rate=calc.annual_rate, years=calc.years,
        final_amount=calc.final_amount, total_contributions=calc.total_contributions, profit=calc.profit,
        schedule=[schemas.YearRow(year=r.year, total=r.total, contributions=r.contributions, profit=r.profit) for r in calc.years_rows],
    )
# ---- DELETE calculation ----
@app.delete("/api/calc/{calc_id}")
def api_delete_calc(
    calc_id: int,
    u: User = Depends(require_user),
    db: Session = Depends(get_db)
):
    calc = (
        db.query(Calculation)
        .filter(
            Calculation.id == calc_id,
            Calculation.user_id == u.id
        )
        .first()
    )

    if not calc:
        raise HTTPException(404, "Расчёт не найден")

    db.delete(calc)
    db.commit()
    return {"ok": True, "deleted_id": calc_id}

# Импорт восстанавливает: title, список годов с (total, contributions, profit), а итоговые поля берёт из последней строки.
# Параметры исходного расчёта (initial_amount, monthly_contribution, annual_rate) восстановить из отчёта нельзя — они сохранятся как 0

@app.get("/stocks", response_class=HTMLResponse)
def stocks_page(request: Request, db: Session = Depends(get_db)):
    return templates.TemplateResponse(
        "stocks.html",
        {
            "request": request,
            "user": get_current_user(request, db)
        }
    )


MOEX_BASE = "https://www.moex.com/iss"


@app.get("/api/moex/{path:path}")
async def moex_proxy(path: str, request: Request):
    """Проксирует любой запрос к iss.moex.com через бэкенд."""
    params = dict(request.query_params)
    url = f"{MOEX_BASE}/{path}.json"
    async with httpx.AsyncClient(timeout=10, follow_redirects=True) as client:
        r = await client.get(url, params=params)
    return Response(
        content=r.content,
        media_type="application/json",
        headers={"Cache-Control": "no-cache"},
    )

@app.get("/api/risk-levels")
def get_risk_levels(db: Session = Depends(get_db)):
    return db.query(RiskLevel).order_by(RiskLevel.id).all()